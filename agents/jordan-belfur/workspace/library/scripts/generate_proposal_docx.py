#!/usr/bin/env python3
"""
Generate CapGrow Partners proposal as a branded Shadow Digital .docx
Branding: Dark headers, red accent (#E8422F), clean sans-serif, professional tables
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# === Brand Colors ===
RED = RGBColor(0xE8, 0x42, 0x2F)       # Shadow red accent
BLACK = RGBColor(0x1A, 0x1A, 0x1A)     # Near-black for body text
DARK = RGBColor(0x2D, 0x2D, 0x2D)      # Dark for headers
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)  # Secondary text
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)  # Table alt rows
TABLE_HEADER_BG = "2D2D2D"              # Dark header bg
TABLE_ALT_BG = "F7F7F7"                 # Alt row bg
RED_HEX = "E8422F"

doc = Document()

# === Page Setup ===
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# === Default Style ===
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = BLACK
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_thin_red_line(doc):
    """Add a thin red horizontal rule"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Create a border-bottom effect using paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{RED_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_section_divider(doc):
    """Add spacing + thin red line between sections"""
    add_thin_red_line(doc)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)


def add_section_heading(doc, number, title):
    """Add a numbered section heading in Shadow style"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    # Number in red
    run_num = p.add_run(f"{number}. ")
    run_num.font.size = Pt(22)
    run_num.font.color.rgb = RED
    run_num.font.name = 'Calibri'
    run_num.bold = True
    # Title in dark
    run_title = p.add_run(title)
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = DARK
    run_title.font.name = 'Calibri'
    run_title.bold = True
    return p


def add_sub_heading(doc, text):
    """Add a sub-heading (bold, slightly larger)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    run.bold = True
    return p


def add_body(doc, text, space_after=Pt(8)):
    """Add body text with optional bold/italic inline formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.3

    # Parse bold and italic markers
    # Split on **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
        else:
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
    return p


def add_bold_lead_body(doc, bold_text, body_text):
    """Add a paragraph with bold lead-in followed by regular text"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(10.5)
    run_bold.font.color.rgb = BLACK
    # Parse remaining text for italic
    parts = re.split(r'(\*.*?\*)', body_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
        else:
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
    return p


def add_bullet(doc, bold_part, rest_text=""):
    """Add a bullet point with optional bold lead"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    if rest_text:
        # Parse for italic
        parts = re.split(r'(\*.*?\*)', rest_text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = BLACK
            else:
                run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = BLACK
    return p


def add_styled_table(doc, headers, rows, col_widths=None, right_align_cols=None):
    """Add a professionally styled table with dark header row"""
    if right_align_cols is None:
        right_align_cols = []

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table style properties
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Remove default borders, we'll style manually
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E8E8E8"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        # Cell padding
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="80" w:type="dxa"/>'
            f'  <w:left w:w="120" w:type="dxa"/>'
            f'  <w:bottom w:w="80" w:type="dxa"/>'
            f'  <w:right w:w="120" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(mar)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            # Alternating row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT

            # Check if text should be bold (last row = total)
            is_bold = row_idx == len(rows) - 1 and (cell_text.startswith("**") or cell_text.startswith("Total") or cell_text.startswith("$"))
            clean_text = cell_text.replace("**", "")

            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
            if is_bold:
                run.bold = True

            # Cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            mar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="60" w:type="dxa"/>'
                f'  <w:left w:w="120" w:type="dxa"/>'
                f'  <w:bottom w:w="60" w:type="dxa"/>'
                f'  <w:right w:w="120" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(mar)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    return table


def add_callout_box(doc, text):
    """Add a callout/note box with left red border"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{RED_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    run.italic = True
    return p


def add_note(doc, text):
    """Add an italicized note"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY_TEXT
    run.italic = True
    return p


# =====================================================================
# COVER PAGE
# =====================================================================

# Add some top spacing
for _ in range(3):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

# "SHADOW" logo text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("SHADOW")
run.font.size = Pt(28)
run.font.name = 'Calibri'
run.font.color.rgb = DARK
run.bold = True
p.paragraph_format.space_after = Pt(4)

# Webflow Enterprise Partner badge text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Webflow Premium Partner  |  Enterprise")
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = GRAY_TEXT
p.paragraph_format.space_after = Pt(36)

# Red accent line
add_thin_red_line(doc)

# Spacer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(12)

# Category + Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("CapGrow Partners LLC  |  Website Redesign & Rebrand")
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RED
run.bold = True
p.paragraph_format.space_after = Pt(8)

# Main headline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Making 20 Years of\nCredibility Visible to the\nPeople Who Need It Most")
run.font.size = Pt(36)
run.font.name = 'Calibri'
run.font.color.rgb = DARK
run.bold = True
p.paragraph_format.space_after = Pt(24)
p.paragraph_format.line_spacing = 1.1

# Intro blurb
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Shadow Digital")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = BLACK
run2 = p.add_run(" is your go-to Webflow agency for building powerful, results-driven websites. We specialize in crafting cutting-edge digital experiences for ")
run2.font.name = 'Calibri'
run2.font.size = Pt(11)
run2.font.color.rgb = BLACK
run3 = p.add_run("modern brands")
run3.bold = True
run3.font.name = 'Calibri'
run3.font.size = Pt(11)
run3.font.color.rgb = BLACK
run4 = p.add_run(", combining sleek design with strategic expertise.")
run4.font.name = 'Calibri'
run4.font.size = Pt(11)
run4.font.color.rgb = BLACK

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(36)
run = p2.add_run("From concept to continuous optimization, we're here to help you achieve your business goals and elevate your web presence.")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = BLACK

# Prepared for / Created by
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Prepared for:")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Sydney Gutierrez")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
run = p.add_run("CapGrow Partners LLC")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Created by:")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Yannick Lorenz")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Shadow Digital")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = BLACK

# === PAGE BREAK ===
doc.add_page_break()


# =====================================================================
# 1. THE OPPORTUNITY
# =====================================================================

add_section_heading(doc, 1, "The Opportunity")

add_body(doc, "CapGrow Partners has spent 20 years building something rare \u2014 a trusted real estate platform that\u2019s helped behavioral health providers across 42 states focus on care instead of property management. Inc. 5000 eight years running. 1,300+ properties. 70+ provider partnerships. The track record speaks for itself.")

add_bold_lead_body(doc, "The problem is, your website doesn\u2019t.", "")

add_body(doc, "Today, CapGrow\u2019s digital presence runs on an aging WordPress site built on a finance-industry template. It wasn\u2019t designed to explain what CapGrow actually does \u2014 and as a result, first-time visitors often walk away confused. Provider executives wonder if CapGrow runs the facilities. PE firms can\u2019t quickly assess the investment model. And smaller operators \u2014 the ones who\u2019d benefit most \u2014 sometimes question whether the offer is legitimate. When your website creates more questions than it answers, every conference handshake has to work twice as hard.")

add_body(doc, "Meanwhile, the competitive landscape has shifted. Both Scioto Properties and Nested have recently redesigned their digital presence, raising the visual bar for the behavioral health real estate category. CapGrow\u2019s business model is different from both \u2014 but online, the distinction is invisible.")

add_body(doc, "This isn\u2019t just a branding exercise. CapGrow\u2019s BD team has tripled in the last year. Den\u00e9, Lexi, and Addie are out at conferences, building relationships, handing out cards \u2014 and every one of those prospects comes home and checks the website. The site is the second impression for every first conversation. Right now, it\u2019s undermining the work your team is doing in the field.")

add_body(doc, "The opportunity: build a digital presence that matches the business CapGrow has already built. A site that instantly communicates the model, projects 20 years of credibility, and gives Sydney the tools to create targeted landing pages for every conference and campaign \u2014 without waiting on developers or fighting a plugin-heavy CMS.")

add_section_divider(doc)


# =====================================================================
# 2. EXPECTED OUTCOMES
# =====================================================================

add_section_heading(doc, 2, "Expected Outcomes")

outcomes = [
    ("Instant clarity for every audience", " \u2014 C-suite executives, non-financial providers, and PE firms will each understand CapGrow\u2019s model within seconds of landing on the site. No more \u201cwait, do you run the houses?\u201d confusion."),
    ("Marketing team independence", " \u2014 Sydney will be able to create, edit, and launch custom landing pages for conferences, campaigns, and partner outreach without writing code or submitting dev tickets. What currently takes weeks will take hours."),
    ("A brand that matches the business", " \u2014 The site will project 20 years of sustained success \u2014 approachable enough for a two-home provider, credible enough for a PE-backed organization evaluating a $100M portfolio deal."),
    ("Competitive differentiation", " \u2014 Visitors will immediately see how CapGrow is different from Scioto (larger, PE-focused) and Nested (younger, tech-forward). The positioning won\u2019t just be different \u2014 it\u2019ll be *clear*."),
    ("A scalable content engine", " \u2014 Blog content, case studies, provider spotlights, and news will live in a structured CMS that\u2019s easy to maintain and built for organic growth. Years of existing blog content preserved and migrated."),
    ("Conference-ready landing pages on demand", " \u2014 Sydney will be able to spin up a branded, targeted landing page for any conference or campaign in hours \u2014 not weeks. No dev tickets, no plugin workarounds, no waiting."),
    ("A site that works as hard as your BD team", " \u2014 Long sales cycles mean prospects come back to the website multiple times before signing. The new site will serve as a nurture tool \u2014 answering objections, building trust, and moving prospects forward between meetings."),
]

for bold_part, rest in outcomes:
    add_bullet(doc, bold_part, rest)

add_section_divider(doc)

# === PAGE BREAK ===
doc.add_page_break()


# =====================================================================
# 3. OUR APPROACH
# =====================================================================

add_section_heading(doc, 3, "Our Approach")

add_body(doc, "We work in gated phases \u2014 each one builds on the last, each one requires your sign-off before we move forward. Nothing gets built that your team hasn\u2019t approved, and nothing gets approved that doesn\u2019t map back to your business goals. Given the strong internal voices on your team, this structure is intentional \u2014 it gives everyone a clear moment to weigh in, react to real visuals, and make decisions before we build further.")

phases = [
    ("Phase 1 \u2014 Align", " \u2014 We start by getting on the same page about who the site is really for and what it needs to say. We\u2019ll map your three audience segments, define the messaging hierarchy, and lock in the sitemap and content strategy. You\u2019ll walk away with a blueprint that the whole team has signed off on before a single pixel gets designed."),
    ("Phase 2 \u2014 Design & Build", " \u2014 You\u2019ll see your core pages come to life in a staging environment early, giving your team something real to react to \u2014 not static mockups that feel abstract. We build in Webflow from the start, so what you\u2019re reviewing is what you\u2019re getting. The component library gets built alongside the pages, so Sydney has a toolkit of pre-approved elements to remix later."),
    ("Phase 3 \u2014 Content & Migration", " \u2014 Blog content, case studies, and institutional pages get migrated with SEO intact. We handle redirect mapping, metadata, and schema markup so nothing gets lost in the transition. New content structures make it easy to keep publishing after launch."),
    ("Phase 4 \u2014 Launch & Enable", " \u2014 Before go-live, we run a full quality pass across browsers and devices. Then we hand you the keys \u2014 with documentation, a hands-on walkthrough, and a dedicated support window so the transition feels seamless, not stressful."),
]

for bold_part, rest in phases:
    add_bold_lead_body(doc, bold_part, rest)

add_body(doc, "We\u2019ll need your team engaged at key approval gates \u2014 typically 2\u20133 touchpoints per phase. Between gates, we execute. You don\u2019t need to manage us.")

add_section_divider(doc)


# =====================================================================
# 4. TIMELINE
# =====================================================================

add_section_heading(doc, 4, "Timeline")

add_bold_lead_body(doc, "Estimated duration: ", "~10\u201312 weeks")

timeline_headers = ["Milestone", "Timing", "What You\u2019ll See"]
timeline_rows = [
    ["Kickoff & Brand Alignment", "Weeks 1\u20132", "Audience mapping finalized, messaging framework locked, sitemap approved"],
    ["Design Concepts & Core Pages", "Weeks 3\u20135", "Homepage and 2\u20133 key pages designed in staging, component library taking shape"],
    ["Full Site Build", "Weeks 6\u20138", "All pages built in Webflow, CMS populated, content migrated"],
    ["QA & Pre-Launch Review", "Week 9", "Full browser/device testing, team walkthrough, final approvals"],
    ["Launch", "Week 10", "Live site, DNS cutover, monitoring active"],
    ["Handoff & Support", "Weeks 11\u201312", "Documentation delivered, hands-on walkthrough, support window open"],
]

add_styled_table(doc, timeline_headers, timeline_rows, col_widths=[2.2, 1.2, 3.6])

add_note(doc, "Timeline assumes timely client feedback at each milestone. Delays in approvals will shift subsequent phases accordingly.")

add_section_divider(doc)

# === PAGE BREAK ===
doc.add_page_break()


# =====================================================================
# 5. INVESTMENT
# =====================================================================

add_section_heading(doc, 5, "Investment")

add_body(doc, "Everything you need for a complete website redesign, rebrand, and Webflow migration \u2014 from strategy through launch. The breakdown below shows exactly where the investment goes, so your team can adjust scope if needed.")

add_sub_heading(doc, "Complete Package \u2014 $66,700")

inv_headers = ["Workstream", "Investment", "What\u2019s Included"]
inv_rows = [
    ["Brand Foundation & UX Strategy", "$9,350", "Brand direction & moodboarding (1 phase), user journey mapping (1), information architecture & sitemap (1)"],
    ["Website Design", "$16,300", "Homepage design (1), core page design (5), supporting page design (4), blog/resource template design (2), custom graphics (12), motion direction (1 phase)"],
    ["Webflow Development", "$21,500", "Webflow project setup (1), component library build (16), homepage development (1), core page development (5), supporting page development (4), Webflow native animations (8)"],
    ["CMS & Content Migration", "$15,500", "CMS collection setup (4), CMS template development (3), content migration (~50 pages), redirect mapping"],
    ["SEO Foundation, Integrations & Launch", "$4,050", "Technical SEO foundation (1), on-page SEO optimization (10 pages), form integration (3), QA & launch (1)"],
    ["**Total**", "**$66,700**", ""],
]

add_styled_table(doc, inv_headers, inv_rows, col_widths=[2.3, 1.0, 3.7], right_align_cols=[1])

add_callout_box(doc, "Every workstream maps to a detailed line-item scope (provided in the accompanying SOW). If your team needs to adjust the investment, we can walk through which items to scale back on our call.")

# Optional Add-Ons
add_sub_heading(doc, "Optional Add-On Services")

add_body(doc, "Additional services that extend the engagement beyond the core redesign. Add any that align with your priorities.")

addon_headers = ["Module", "Investment", "What It Covers"]
addon_rows = [
    ["Training & Team Enablement", "$1,750", "Dedicated enablement program: live CMS training (60\u201390 min), website editing training (60 min), marketing team enablement workshop (60\u201390 min). Recommended for teams where one person manages the site."],
    ["SEO Audit & Keyword Strategy", "$2,100", "Keyword research for behavioral health real estate niche, competitor SEO analysis, content gap identification, keyword-to-page mapping"],
    ["WAIO Implementation (AI Discoverability)", "$5,000", "Shadow Digital\u2019s proprietary AEO framework \u2014 ensures CapGrow surfaces accurately across ChatGPT, Perplexity, and Gemini. Includes FAQ schema, entity alignment, and semantic structure across the full site."],
    ["Content Outlines", "$150/outline", "SEO-driven content outline per page \u2014 heading structure, topic guidance, keyword direction, audience-specific messaging recommendations"],
    ["Content Writing", "$350/page", "Full copywriting per page based on approved outline \u2014 audience-tailored messaging, SEO-optimized, ready to publish"],
]

add_styled_table(doc, addon_headers, addon_rows, col_widths=[2.3, 1.0, 3.7], right_align_cols=[1])

# Included bonuses
add_sub_heading(doc, "Included at No Additional Cost")

add_bullet(doc, "Post-Launch Support \u2014 2 Weeks", " (Valued at $4,200) \u2014 Full team availability for bug fixes, content adjustments, and transition support during the critical post-launch window. We don\u2019t disappear after launch \u2014 we stay until your team is fully comfortable.")

add_bullet(doc, "Redirect Mapping & SEO Preservation", " (Valued at $2,100) \u2014 Full URL-by-URL redirect mapping and validation so you don\u2019t lose a single page of organic equity during the migration. Every URL accounted for, every redirect tested.")

add_body(doc, "That\u2019s **$6,300 in additional value** included with this engagement.")

# Payment Structure
add_sub_heading(doc, "Payment Structure")

pay_headers = ["Milestone", "Percentage", "Amount"]
pay_rows = [
    ["At signing", "50%", "$33,350"],
    ["Net-30 post kickoff", "40%", "$26,680"],
    ["At project handoff", "10%", "$6,670"],
]

add_styled_table(doc, pay_headers, pay_rows, col_widths=[2.5, 1.5, 2.0], right_align_cols=[2])

add_section_divider(doc)

# === PAGE BREAK ===
doc.add_page_break()


# =====================================================================
# 6. WHY SHADOW DIGITAL
# =====================================================================

add_section_heading(doc, 6, "Why Shadow Digital")

why_items = [
    ("WordPress-to-Webflow is our core competency", " \u2014 We\u2019ve executed 50+ enterprise migrations from WordPress, HubSpot, and custom platforms to Webflow. This isn\u2019t a side service \u2014 it\u2019s the foundation of our practice. We know where the bodies are buried with plugin-heavy WordPress sites, and plugin-heavy is an understatement for what\u2019s running on capgrowpartners.com today."),
    ("We\u2019ve solved the complex messaging problem before", " \u2014 Multi-audience companies with misunderstood business models are our sweet spot. We migrated Mizuho Bank\u2019s entire global web presence to Webflow \u2014 a regulated financial institution with similar complexity around audience segmentation, credibility signaling, and regulatory nuance. Their marketing team went from a multi-week content cycle to same-day publishing."),
    ("We haven\u2019t worked with a REIT \u2014 but we\u2019ve worked in your world.", " Our experience spans financial services, fintech, and PE-backed companies where trust, compliance, and institutional credibility are non-negotiable. The challenges CapGrow faces \u2014 explaining a complex model to multiple audiences, projecting stability without feeling corporate, balancing accessibility with sophistication \u2014 are the exact problems we solve."),
    ("Enablement-first, not dependency-first", " \u2014 Every engagement ends with your team trained and independent. We don\u2019t build sites that require you to call us every time you need a new page. The component library and documentation ensure your marketing team can operate autonomously post-launch."),
    ("Built for stakeholder-driven design", " \u2014 We\u2019ve worked with leadership teams that have strong visual opinions and high expectations. Our gated approval process gives every stakeholder clear touchpoints to weigh in on real pages \u2014 not abstract wireframes \u2014 so feedback stays grounded, productive, and doesn\u2019t loop."),
    ("Full-stack when you need it", " \u2014 Shadow Digital is part of a broader network covering brand, digital experience, SEO, performance marketing, and strategy. You get boutique-level focus on your project with enterprise-level depth behind it. If CapGrow\u2019s needs expand into content strategy, paid media, or ongoing SEO, the team is already here."),
]

for bold_part, rest in why_items:
    add_bold_lead_body(doc, bold_part, rest)

add_section_divider(doc)


# =====================================================================
# 7. SELECT CASE STUDIES
# =====================================================================

add_section_heading(doc, 7, "Select Case Studies")

# Bench
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Bench")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK
run2 = p.add_run("  \u2014  ")
run2.font.name = 'Calibri'
run2.font.size = Pt(12)
run2.font.color.rgb = GRAY_TEXT
run3 = p.add_run("Fintech  |  WordPress \u2192 Webflow Migration")
run3.font.name = 'Calibri'
run3.font.size = Pt(10.5)
run3.font.color.rgb = GRAY_TEXT
run3.italic = True

add_body(doc, "Bench came to us with 1,000+ content pieces trapped in a legacy WordPress CMS that their marketing team couldn\u2019t manage independently. We migrated the full library, rebuilt their content architecture in Webflow, and gave their team the tools to publish without engineering support. Within 6 months of launch: 450 new pages created by the marketing team alone, organic keywords grew from 2,000 to 4,900, and the site hit a performance score of 96. Page creation capacity tripled \u2014 the team went from waiting on dev tickets to publishing same-day.")

# TSIA
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("TSIA")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK
run2 = p.add_run("  \u2014  ")
run2.font.name = 'Calibri'
run2.font.size = Pt(12)
run2.font.color.rgb = GRAY_TEXT
run3 = p.add_run("Technology Services  |  Enterprise Webflow Migration")
run3.font.name = 'Calibri'
run3.font.size = Pt(10.5)
run3.font.color.rgb = GRAY_TEXT
run3.italic = True

add_body(doc, "TSIA needed 200+ pages migrated from a legacy platform to Webflow \u2014 without losing search equity or breaking member-facing workflows. We completed the full migration in 6 weeks, improved mobile page speed by 112%, and preserved their entire SEO footprint through redirect mapping and metadata migration. Their team called it \u201can outstanding performance \u2014 delivered on time, on budget, and beyond expectations.\u201d")

# Sterling Bank
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Sterling Bank")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK
run2 = p.add_run("  \u2014  ")
run2.font.name = 'Calibri'
run2.font.size = Pt(12)
run2.font.color.rgb = GRAY_TEXT
run3 = p.add_run("Financial Services  |  WordPress \u2192 Webflow")
run3.font.name = 'Calibri'
run3.font.size = Pt(10.5)
run3.font.color.rgb = GRAY_TEXT
run3.italic = True

add_body(doc, "Sterling Bank\u2019s marketing team was dependent on their engineering department for every website update \u2014 a bottleneck that slowed campaigns and limited their ability to respond to market conditions. We migrated their site to Webflow with a component library designed for marketing independence. Post-launch, their team manages the site end-to-end without developer involvement \u2014 the same outcome we\u2019re designing for Sydney and the CapGrow marketing function.")

add_section_divider(doc)

# === PAGE BREAK ===
doc.add_page_break()


# =====================================================================
# 8. NEXT STEPS
# =====================================================================

add_section_heading(doc, 8, "Next Steps")

steps = [
    ("1. Review this proposal", " with your team"),
    ("2. Walkthrough call", " \u2014 We\u2019ll discuss questions, refine scope, and finalize the package (scheduled: February 26)"),
    ("3. SOW delivery", " \u2014 We\u2019ll send the detailed Statement of Work with deliverables, hours, and terms based on what we align on"),
    ("4. Sign & kick off", " \u2014 Lock in the engagement and schedule kickoff within 1 week of signing"),
]

for bold_part, rest in steps:
    add_bold_lead_body(doc, bold_part, rest)

# Spacer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(8)

add_body(doc, "Your BD team is actively driving traffic to a site that doesn\u2019t convert \u2014 every week without a new presence is a week of lost follow-through from conferences and prospect conversations. With Scioto and Nested having already redesigned, the window to lead the category is now.")

add_note(doc, "Post-launch, we also offer ongoing support retainers for continued optimization, content updates, and campaign support \u2014 happy to discuss once the site is live.")


# =====================================================================
# SAVE
# =====================================================================

output_path = "/Users/yaya/Claude/van-agent-sdk/agents/jordan-belfur/workspace/deals/capgrow-partners/deliverables/CapGrow Partners - Website Redesign Proposal.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
