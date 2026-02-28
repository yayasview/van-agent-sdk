#!/usr/bin/env python3
"""
Generate Sofar Ocean proposal as a branded Shadow Digital .docx
Two-option structure: Dev Only ($56,100) vs Full Service ($78,000)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# === Brand Colors ===
RED = RGBColor(0xE8, 0x42, 0x2F)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)
TABLE_HEADER_BG = "2D2D2D"
TABLE_ALT_BG = "F7F7F7"
RED_HEX = "E8422F"

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_thin_red_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{RED_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_section_divider(doc):
    add_thin_red_line(doc)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run_num = p.add_run(f"{number}. ")
    run_num.font.size = Pt(22)
    run_num.font.color.rgb = RED
    run_num.font.name = 'Calibri'
    run_num.bold = True
    run_title = p.add_run(title)
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = DARK
    run_title.font.name = 'Calibri'
    run_title.bold = True
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    run.bold = True
    return p


def add_body(doc, text, space_after=Pt(8)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.3
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    return p


def add_bold_lead_body(doc, bold_text, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(10.5)
    run_bold.font.color.rgb = BLACK
    parts = re.split(r'(\*.*?\*)', body_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    return p


def add_bullet(doc, bold_part, rest_text=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
    if rest_text:
        parts = re.split(r'(\*.*?\*)', rest_text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = BLACK
    return p


def add_styled_table(doc, headers, rows, col_widths=None, right_align_cols=None):
    if right_align_cols is None:
        right_align_cols = []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E8E8E8"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>')
        tcPr.append(mar)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
            is_bold = row_idx == len(rows) - 1 and ("**" in cell_text or cell_text.startswith("Total") or cell_text.startswith("$"))
            clean_text = cell_text.replace("**", "")
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
            if is_bold:
                run.bold = True
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="60" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:bottom w:w="60" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>')
            tcPr.append(mar)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)
    return table


def add_callout_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="{RED_HEX}"/></w:pBdr>')
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    run.italic = True
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY_TEXT
    run.italic = True
    return p


# =====================================================================
# COVER PAGE
# =====================================================================

for _ in range(3):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("SHADOW")
run.font.size = Pt(28)
run.font.name = 'Calibri'
run.font.color.rgb = DARK
run.bold = True
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Webflow Premium Partner  |  Enterprise")
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = GRAY_TEXT
p.paragraph_format.space_after = Pt(36)

add_thin_red_line(doc)

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Sofar Ocean  |  Website Rebuild")
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RED
run.bold = True
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Giving the World\u2019s Ocean\nData Platform a Digital\nPresence That Moves\nLike Water")
run.font.size = Pt(36)
run.font.name = 'Calibri'
run.font.color.rgb = DARK
run.bold = True
p.paragraph_format.space_after = Pt(24)
p.paragraph_format.line_spacing = 1.1

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Shadow Digital")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = BLACK
run2 = p.add_run(" is your go-to Webflow agency for building powerful, results-driven websites. We specialize in crafting cutting-edge digital experiences for ")
run2.font.name = 'Calibri'
run2.font.size = Pt(11)
run2.font.color.rgb = BLACK
run3 = p.add_run("modern brands")
run3.bold = True
run3.font.name = 'Calibri'
run3.font.size = Pt(11)
run3.font.color.rgb = BLACK
run4 = p.add_run(", combining sleek design with strategic expertise.")
run4.font.name = 'Calibri'
run4.font.size = Pt(11)
run4.font.color.rgb = BLACK

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(36)
run = p2.add_run("From concept to continuous optimization, we\u2019re here to help you achieve your business goals and elevate your web presence.")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = BLACK

# Prepared for / Created by
for label, name, company in [
    ("Prepared for:", "Rosy Garcia", "Sofar Ocean"),
    ("Created by:", "Yannick Lorenz", "Shadow Digital"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(name)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(company)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

doc.add_page_break()


# =====================================================================
# 1. THE OPPORTUNITY
# =====================================================================

add_section_heading(doc, 1, "The Opportunity")

add_body(doc, "Sofar Ocean has built something extraordinary \u2014 the world\u2019s largest real-time ocean sensor network, powering decisions across maritime shipping, climate science, defense, and coastal resilience. Spotter buoys floating in every ocean. Wayfinder routing ships around weather. Forecasting models that outperform legacy providers. The technology speaks for itself.")

add_bold_lead_body(doc, "The website doesn\u2019t.", "")

add_body(doc, "What you have today is what your team calls a \u201clasagna\u201d \u2014 layers of additions built over years without a unifying architecture, component system, or content strategy. Every page was built from scratch. Changes in one section break things in another. The marketing team is afraid to touch it. And when a prospect lands on sofarocean.com after a conference conversation or a Wayfinder demo, they\u2019re met with a site that doesn\u2019t reflect the sophistication of the product behind it.")

add_body(doc, "This matters more now than it did a year ago. Sofar\u2019s audience has split into two distinct segments \u2014 digitally native buyers who expect a modern, conversion-driven experience, and traditional maritime operators who need a clear, credible overview they can share with their fleet management teams. The current site serves neither well. There\u2019s no solutions-based navigation, no persona-specific entry points, and no way for Rosy or Michelle to create targeted content without rebuilding pages from scratch.")

add_body(doc, "Meanwhile, the team has grown. Rosy, Michelle, and Maggie are the first dedicated marketing and design hires focused on the website \u2014 but they\u2019ve inherited a platform with no documentation, no component library, and no templates. The tools don\u2019t match the talent.")

add_body(doc, "The opportunity: a clean-slate rebuild that gives Sofar a site architecture worthy of the technology it represents \u2014 solutions-based navigation, motion-heavy design that evokes the ocean, a component library the team can actually use, and a CMS structure that scales with 200+ pieces of existing content and everything that comes next.")

add_section_divider(doc)


# =====================================================================
# 2. EXPECTED OUTCOMES
# =====================================================================

add_section_heading(doc, 2, "Expected Outcomes")

outcomes = [
    ("Solutions-based navigation that sells", " \u2014 Visitors will find what they need through the lens of what Sofar solves (voyage intelligence, ocean sensing, weather forecasting, defense) \u2014 not through a product catalog. Each solution area will have its own entry point, use cases, and conversion path."),
    ("Two audiences, one site", " \u2014 Digitally native buyers get the interactive, conversion-optimized experience they expect. Traditional maritime operators get the clean, authoritative overview they need to share internally. Both audiences self-select within seconds of landing."),
    ("A site that moves like the ocean", " \u2014 Parallax scrolling, native Webflow animations, and motion design throughout \u2014 not as decoration, but as a storytelling device that reinforces Sofar\u2019s identity. The site will feel alive in a way that matches the real-time data flowing through your sensor network."),
    ("Marketing team independence", " \u2014 Rosy, Michelle, and Maggie will be able to create, edit, and launch pages using a documented component library and pre-built templates \u2014 without rebuilding from scratch every time. What currently requires starting from zero will be drag-and-drop."),
    ("A scalable content engine", " \u2014 200+ blog posts, research papers, case studies, and event listings migrated into a structured CMS with proper taxonomy, filtering, and templates. New content types can be added without developer involvement."),
    ("Clean architecture, zero tech debt", " \u2014 Every component named, documented, and reusable. No more \u201clasagna\u201d \u2014 a system where changes in one place don\u2019t break things elsewhere. The team won\u2019t be afraid to touch the site anymore."),
    ("SEO foundation from day one", " \u2014 Technical SEO, on-page optimization, redirect mapping, and schema markup implemented during the build \u2014 not bolted on afterward. The organic equity you\u2019ve built over years gets preserved and strengthened."),
]

for bold_part, rest in outcomes:
    add_bullet(doc, bold_part, rest)

add_section_divider(doc)
doc.add_page_break()


# =====================================================================
# 3. OUR APPROACH
# =====================================================================

add_section_heading(doc, 3, "Our Approach")

add_body(doc, "We work in sprint-based cycles \u2014 each sprint delivers tangible progress, each major phase requires your sign-off before we move forward. Nothing gets built that your team hasn\u2019t reviewed, and nothing gets approved that doesn\u2019t map back to your business goals.")

add_bold_lead_body(doc, "Design handoff model (Dev Only):", " Maggie delivers Figma designs sprint-by-sprint, and our team translates them into production Webflow \u2014 with a design consultant embedded in the first two weeks to ensure every design decision is Webflow-feasible before a single pixel gets built. This keeps the creative vision in-house while our team handles the technical execution.")

add_bold_lead_body(doc, "Full service model:", " We handle everything \u2014 discovery, information architecture, UX wireframes, UI design, motion direction, and development. Your team reviews at structured gates throughout, and Maggie stays involved as a design collaborator rather than the sole design owner.")

add_body(doc, "**Regardless of path, every engagement includes:**")

phases = [
    ("Phase 1 \u2014 Foundation", " \u2014 We align on the solutions-based sitemap, lock in the component architecture, and run an SEO audit so we\u2019re building on solid ground. In the dev-only path, this is where our design consultant reviews Maggie\u2019s approach and flags anything that won\u2019t translate cleanly to Webflow."),
    ("Phase 2 \u2014 Build", " \u2014 Core pages come to life in a staging environment. The component library gets built alongside the pages, so your team has reusable building blocks from day one. Motion and interaction design are woven into the build \u2014 not added as an afterthought."),
    ("Phase 3 \u2014 Content & Migration", " \u2014 200+ pieces of content migrated to the new CMS structure with proper taxonomy, filtering, and templates. Blog posts, research papers, case studies, and event listings all get structured homes. Redirect mapping ensures nothing gets lost."),
    ("Phase 4 \u2014 Launch & Enable", " \u2014 Full QA across browsers and devices, SEO validation, and then a comprehensive handoff \u2014 documentation, component walkthrough, and hands-on training so the team is self-sufficient from day one."),
]

for bold_part, rest in phases:
    add_bold_lead_body(doc, bold_part, rest)

add_body(doc, "We\u2019ll need your team engaged at key review points \u2014 typically once per sprint. Between reviews, we execute. You don\u2019t need to manage us.")

add_section_divider(doc)


# =====================================================================
# 4. TIMELINE
# =====================================================================

add_section_heading(doc, 4, "Timeline")

add_sub_heading(doc, "Option A \u2014 Dev Only: ~14 weeks (12 build + 2 support)")

tl_a_headers = ["Milestone", "Timing", "What You\u2019ll See"]
tl_a_rows = [
    ["Project Setup & Design Consulting", "Weeks 1\u20132", "SEO audit complete, design feasibility review, component architecture locked"],
    ["Core Development Begins", "Weeks 3\u20134", "Homepage and priority pages in staging, motion/interaction work starts"],
    ["Full Build Velocity", "Weeks 5\u20138", "All solution pages built, CMS populated, content migration underway"],
    ["Content Migration & SEO", "Weeks 9\u201310", "200+ posts migrated, redirects mapped, on-page SEO implemented"],
    ["QA, Enablement & Launch", "Weeks 11\u201312", "Full browser/device testing, team training, go-live"],
    ["Post-Launch Support", "Weeks 13\u201314", "Bug fixes, content adjustments, transition support"],
]
add_styled_table(doc, tl_a_headers, tl_a_rows, col_widths=[2.2, 1.2, 3.6])

add_sub_heading(doc, "Option B \u2014 Full Service: ~20 weeks (18 build + 2 support)")

tl_b_headers = ["Milestone", "Timing", "What You\u2019ll See"]
tl_b_rows = [
    ["Discovery & Information Architecture", "Weeks 1\u20132", "Solutions-based sitemap finalized, audience mapping complete"],
    ["UX Wireframes", "Weeks 3\u20134", "Key page structures approved, design review"],
    ["Visual Design & Motion Direction", "Weeks 5\u20138", "Full UI design, motion concepts, design handoff to dev"],
    ["Project Setup & Core Dev", "Weeks 9\u201310", "SEO audit, homepage and priority pages in staging"],
    ["Full Build Velocity", "Weeks 11\u201314", "All pages built, motion/interaction implemented, content migration"],
    ["QA, Enablement & Launch", "Weeks 15\u201318", "Full QA, SEO validation, team training, go-live"],
    ["Post-Launch Support", "Weeks 19\u201320", "Bug fixes, adjustments, transition support"],
]
add_styled_table(doc, tl_b_headers, tl_b_rows, col_widths=[2.2, 1.2, 3.6])

add_note(doc, "Both timelines assume designs are delivered on schedule (dev-only) or timely client feedback at each milestone (full service). Delays shift subsequent phases accordingly.")

add_section_divider(doc)
doc.add_page_break()


# =====================================================================
# 5. INVESTMENT
# =====================================================================

add_section_heading(doc, 5, "Investment")

add_body(doc, "Two options, depending on whether Sofar\u2019s in-house design team leads the creative direction or Shadow Digital handles design end-to-end.")

# Option A
add_sub_heading(doc, "Option A \u2014 Dev Only: $56,100")

add_body(doc, "Maggie leads design in Figma. Shadow Digital translates designs into production Webflow, handles all development, motion, migration, SEO, and enablement. Includes design consulting to ensure Webflow feasibility.")

inv_a_headers = ["Work Stream", "Hours", "Investment"]
inv_a_rows = [
    ["Project Management", "56", "$8,400"],
    ["Webflow Development", "158", "$23,700"],
    ["Motion & Interaction Development", "48", "$7,200"],
    ["Content Migration", "40", "$6,000"],
    ["SEO Setup & Optimization", "32", "$4,800"],
    ["Quality Assurance", "24", "$3,600"],
    ["Design Consulting", "16", "$2,400"],
    ["**Total**", "**374**", "**$56,100**"],
]
add_styled_table(doc, inv_a_headers, inv_a_rows, col_widths=[3.0, 1.0, 1.5], right_align_cols=[1, 2])

# Option B
add_sub_heading(doc, "Option B \u2014 Full Service: $78,000")

add_body(doc, "Shadow Digital handles everything \u2014 discovery, IA, UX, UI, motion direction, development, migration, SEO, QA, and enablement.")

inv_b_headers = ["Work Stream", "Hours", "Investment"]
inv_b_rows = [
    ["Project Management", "72", "$10,800"],
    ["UX Design", "80", "$12,000"],
    ["UI Design", "76", "$11,400"],
    ["Creative Direction", "24", "$3,600"],
    ["Webflow Development", "136", "$20,400"],
    ["Motion & Interaction Development", "48", "$7,200"],
    ["Content Migration", "40", "$6,000"],
    ["SEO Setup & Optimization", "24", "$3,600"],
    ["Quality Assurance", "20", "$3,000"],
    ["**Total**", "**520**", "**$78,000**"],
]
add_styled_table(doc, inv_b_headers, inv_b_rows, col_widths=[3.0, 1.0, 1.5], right_align_cols=[1, 2])

# Comparison
add_sub_heading(doc, "What\u2019s the Difference?")

diff_headers = ["", "Dev Only ($56,100)", "Full Service ($78,000)"]
diff_rows = [
    ["Who designs", "Maggie (in-house) + SD consulting", "Shadow Digital end-to-end"],
    ["Timeline", "12 weeks + support", "18 weeks + support"],
    ["Design risk", "Maggie is new, desktop-only Figma \u2014 SD consults on feasibility", "Full design team with Webflow-native expertise"],
    ["Motion", "SD builds from Maggie\u2019s direction", "SD handles motion direction + execution"],
    ["Best for", "Strong internal design vision, budget priority", "Need a partner to own the full creative process"],
]
add_styled_table(doc, diff_headers, diff_rows, col_widths=[1.3, 2.6, 2.6])

add_callout_box(doc, "Both options include the same scope of development, content migration, SEO, and enablement. The difference is who owns the design process.")

# Bonuses
add_sub_heading(doc, "Included at No Additional Cost")

add_bullet(doc, "Post-Launch Support \u2014 2 Weeks", " (Valued at $4,200) \u2014 Full team availability for bug fixes, content adjustments, and transition support during the critical post-launch window.")
add_bullet(doc, "Redirect Mapping & SEO Preservation", " (Valued at $2,100) \u2014 Full URL-by-URL redirect mapping so you don\u2019t lose organic equity during the rebuild.")
add_bullet(doc, "Component Library Documentation", " (Valued at $1,500) \u2014 Every component named, documented, and annotated so the team can build independently post-launch.")

add_body(doc, "That\u2019s **$7,800 in additional value** included with either option.")

# Payment
add_sub_heading(doc, "Payment Structure")

pay_headers = ["Milestone", "Percentage", "Option A", "Option B"]
pay_rows = [
    ["At signing", "50%", "$28,050", "$39,000"],
    ["Net-30 post kickoff", "40%", "$22,440", "$31,200"],
    ["At project handoff", "10%", "$5,610", "$7,800"],
]
add_styled_table(doc, pay_headers, pay_rows, col_widths=[2.0, 1.0, 1.5, 1.5], right_align_cols=[2, 3])

add_section_divider(doc)
doc.add_page_break()


# =====================================================================
# 6. WHY SHADOW DIGITAL
# =====================================================================

add_section_heading(doc, 6, "Why Shadow Digital")

why_items = [
    ("Motion-forward Webflow is what we do", " \u2014 We\u2019ve built 50+ enterprise sites on Webflow with native animations, parallax, and interaction design baked in from the start \u2014 not bolted on afterward. The ocean-inspired motion Sofar wants isn\u2019t a stretch for us \u2014 it\u2019s our standard."),
    ("We\u2019ve rebuilt \u201clasagna\u201d sites before", " \u2014 Companies that have layered additions for years without architecture are our sweet spot. We don\u2019t just redesign \u2014 we create a system: named components, documented patterns, and clean class structures so the next person who touches the site understands exactly what they\u2019re working with."),
    ("Enablement-first, not dependency-first", " \u2014 Your team is building Webflow expertise from scratch. Every engagement ends with documentation, training, and a component library designed for independence. We don\u2019t create dependency \u2014 we build capability. Six months post-launch, your team should be shipping pages without thinking about us."),
    ("We\u2019ve handled complex content migrations at scale", " \u2014 TSIA: 200+ pages migrated in 6 weeks with 112% mobile speed improvement and zero SEO loss. Bench: 1,000+ content pieces migrated, marketing team creating 450 new pages within 6 months. Your 200+ blog posts, research papers, and case studies are well within our operating range."),
    ("Sprint-based delivery with real accountability", " \u2014 You\u2019ll see progress every two weeks in a staging environment. Not mockups, not wireframes \u2014 real pages you can click through. If something\u2019s off, we catch it in Week 3, not Week 12."),
]

for bold_part, rest in why_items:
    add_bold_lead_body(doc, bold_part, rest)

add_section_divider(doc)


# =====================================================================
# 7. SELECT CASE STUDIES
# =====================================================================

add_section_heading(doc, 7, "Select Case Studies")

# Bench
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Bench")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK
run2 = p.add_run("  \u2014  ")
run2.font.name = 'Calibri'
run2.font.size = Pt(12)
run2.font.color.rgb = GRAY_TEXT
run3 = p.add_run("Fintech  |  WordPress \u2192 Webflow Migration")
run3.font.name = 'Calibri'
run3.font.size = Pt(10.5)
run3.font.color.rgb = GRAY_TEXT
run3.italic = True

add_body(doc, "Bench came to us with 1,000+ content pieces trapped in a legacy CMS that their marketing team couldn\u2019t manage. We rebuilt their content architecture in Webflow and gave their team full publishing independence. Within 6 months: 450 new pages created by the marketing team alone, organic keywords grew from 2,000 to 4,900, and page creation capacity tripled. The same outcome we\u2019re designing for Rosy, Michelle, and Maggie.")

# TSIA
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("TSIA")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK
run2 = p.add_run("  \u2014  ")
run2.font.name = 'Calibri'
run2.font.size = Pt(12)
run2.font.color.rgb = GRAY_TEXT
run3 = p.add_run("Technology Services  |  Enterprise Webflow Migration")
run3.font.name = 'Calibri'
run3.font.size = Pt(10.5)
run3.font.color.rgb = GRAY_TEXT
run3.italic = True

add_body(doc, "TSIA needed 200+ pages migrated to Webflow without losing search equity \u2014 a scale nearly identical to Sofar\u2019s content library. We completed the full migration in 6 weeks, improved mobile page speed by 112%, and preserved their entire SEO footprint. Their team called it \u201can outstanding performance \u2014 delivered on time, on budget, and beyond expectations.\u201d")

add_section_divider(doc)
doc.add_page_break()


# =====================================================================
# 8. NEXT STEPS
# =====================================================================

add_section_heading(doc, 8, "Next Steps")

steps = [
    ("1. Review both options", " with your team"),
    ("2. Alignment call", " \u2014 We\u2019ll walk through the proposal, answer questions, and help you decide between dev-only and full service"),
    ("3. SOW delivery", " \u2014 Detailed Statement of Work based on the option you select"),
    ("4. Sign & kick off", " \u2014 Lock in the engagement and schedule kickoff for the week of March 3"),
]

for bold_part, rest in steps:
    add_bold_lead_body(doc, bold_part, rest)

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(8)

add_body(doc, "Sofar\u2019s website is the front door for every prospect who\u2019s seen a Spotter demo, attended a webinar, or met your team at a conference. Right now, the site behind that front door doesn\u2019t match the product behind it. The team is in place. The sitemap is ready. The only thing missing is the build partner.")

add_body(doc, "We\u2019re ready to start when you are. Yannick is available at yaya@vezadigital.com to discuss anything before the alignment call.")

add_note(doc, "Post-launch, we also offer ongoing support retainers for continued optimization, content updates, and campaign support \u2014 happy to discuss once the site is live.")


# =====================================================================
# SAVE
# =====================================================================

output_path = "/Users/yaya/Claude/van-agent-sdk/agents/jordan-belfur/workspace/deals/sofar-ocean/deliverables/Sofar Ocean - Website Rebuild Proposal.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
